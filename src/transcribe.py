"""
Urdu Speech-to-Text Transcription Tool  v2
==========================================
Model: kingabzpro/whisper-large-v3-turbo-urdu
Fine-tuned on: openai/whisper-large-v3-turbo (Common Voice 17.0 - Urdu)

Configuration:
  All settings are loaded from .env in the project root.
  See .env for available options and defaults.

Usage (single file):
  python transcribe.py trans audio.mp3
  python transcribe.py trans audio.wav --output result.txt
  python transcribe.py trans audio.mp3 --format docx
  python transcribe.py trans audio.mp3 --quiet --language en
  python transcribe.py trans --dry-run audio.mp3          # preview only
  python transcribe.py trans audio.mp3 --spell             # NEW: Apply Urdu spell correction

Usage (batch):
  python transcribe.py batch ./folder/                    # process all files
  python transcribe.py batch ./folder/ -w 4               # 4 parallel workers
  python transcribe.py batch ./folder/ --no-parallel       # force sequential
  python transcribe.py batch --dry-run-batch ./folder/    # preview only

Usage (config):
  python transcribe.py config show                        # view settings
  python transcribe.py config set model=my-model language=en

Usage (help):
  python transcribe.py --help                             # full help
  python transcribe.py trans --help                       # single-file help
  python transcribe.py batch --help                       # batch help
"""

import sys
import os
import argparse
import json
import hashlib
import re
import io
import time
from pathlib import Path
from datetime import datetime, timezone
from concurrent.futures import ProcessPoolExecutor, as_completed


# ════════════════════════════════════════════════════════════════════
# Load .env configuration (reads project/.env if present)
# ════════════════════════════════════════════════════════════════════

def _env_str(key: str, default: str) -> str:
    return os.environ.get(key, default)

def _env_int(key: str, default: int) -> int:
    try:
        return int(os.environ.get(key, str(default)))
    except (ValueError, TypeError):
        return default

def _env_float(key: str, default: float) -> float:
    try:
        return float(os.environ.get(key, str(default)))
    except (ValueError, TypeError):
        return default

def _env_bool(key: str, default: bool) -> bool:
    val = os.environ.get(key)
    if val is None:
        return default
    return val.lower() in ("true", "1", "yes")

# ── Load .env file from project root if it exists ──────────────────
_env_file = Path(__file__).parents[1] / ".env"
if _env_file.exists():
    with open(_env_file, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, _, value = line.partition("=")
            key = key.strip()
            value = value.strip()
            if key and key not in os.environ:
                # Preserve type hints for known keys
                numeric_keys = {"CHUNK_DURATION_S", "CHUNK_OVERLAP_S", "MAX_AUDIO_BEFORE_SPLIT_S",
                                                "MIN_AUDIO_DURATION_S", "DEFAULT_WORKERS", "MAX_SEGMENT_MINUTES",
                                                "YOUTUBE_MIN_DURATION_S", "REPEAT_THRESHOLD",
                                                "STRIDE_MIN_OVERLAP_S", "SHORT_AUDIO_MAX_SEGMENT_SRT",
                                                "SHORT_AUDIO_MAX_SEGMENT_TXT", "MAX_WORKERS", "WORKERS_PER_CPU"}
                if key in numeric_keys:
                    try:
                        os.environ[key] = str(float(value))
                    except ValueError:
                        pass
                else:
                    os.environ[key] = value
    del _env_file


# ── Optional dependencies with graceful fallbacks ─────────────────────────────

try:
    from tqdm import tqdm
    HAS_TQDM = True
except ImportError:
    HAS_TQDM = False

    class tqdm:  # type: ignore
        """Minimal tqdm fallback."""
        def __init__(self, *args, **kwargs):
            self._disable = kwargs.get("disable", False)
            self._desc = kwargs.get("desc", "")
            self._total = kwargs.get("total", 0)
            self._n = 0
        def __enter__(self):
            return self
        def __exit__(self, *args):
            pass
        @staticmethod
        def update(*a):
            pass
        def set_description(self, desc):
            self._desc = desc
        def refresh(self):
            pass


# ── Constants ─────────────────────────────────────────────────────────────────

SUPPORTED_EXTENSIONS = set(ext.strip() for ext in _env_str("SUPPORTED_EXTENSIONS", ".mp3,.wav,.flac,.ogg,.m4a,.mp4,.webm,.aac,.wma").split(","))

CACHE_FILE = _env_str("CACHE_FILE", ".transcribe_cache.json")
CONFIG_FILE = _env_str("CONFIG_FILE", "transcribe_config.json")
BATCH_STATUS_FILE = _env_str("BATCH_STATUS_FILE", ".batch_status.json")

DEFAULT_MODEL = _env_str("HF_MODEL_ID", "kingabzpro/whisper-large-v3-turbo-urdu")

# ── Low-end system detection (RAM-aware) ─────────────────────────────────

try:
    import ctypes
    kernel32 = ctypes.windll.kernel32
    class _MEMORYSTATUSEX(ctypes.Structure):
        _fields_ = [
            ("dwLength", ctypes.c_ulong),
            ("dwMemoryLoad", ctypes.c_ulong),
            ("ullTotalPhys", ctypes.c_ulonglong),
            ("ullAvailPhys", ctypes.c_ulonglong),
            ("ullTotalPageFile", ctypes.c_ulonglong),
            ("ullAvailPageFile", ctypes.c_ulonglong),
            ("ullTotalVirtual", ctypes.c_ulonglong),
            ("ullAvailVirtual", ctypes.c_ulonglong),
            ("ullAvailExtendedVirtual", ctypes.c_ulonglong),
        ]
    def _get_ram_windows():
        stat = _MEMORYSTATUSEX()
        stat.dwLength = ctypes.sizeof(_MEMORYSTATUSEX)
        if kernel32.GlobalMemoryStatusEx(ctypes.byref(stat)):
            return stat.ullTotalPhys
    _TOTAL_RAM_BYTES = _get_ram_windows() or 0
except Exception:
    _TOTAL_RAM_BYTES = 0

def _get_system_ram_mb():
    """Get total system RAM in MB. Falls back to 16384 if detection fails."""
    ram_mb = _TOTAL_RAM_BYTES // (1024 * 1024) if _TOTAL_RAM_BYTES > 0 else None
    fallback = int(os.environ.get("SYSTEM_RAM_MB", "16384"))
    return ram_mb or fallback

_SYSTEM_RAM_MB = _get_system_ram_mb()
_LOW_RAM_MODE = os.environ.get("LOW_RAM", "auto").lower() in ("auto", "true", "1") and _SYSTEM_RAM_MB < 8192
if _LOW_RAM_MODE:
    print(f"[INFO] Low-RAM mode detected ({_SYSTEM_RAM_MB} MB) — optimizations enabled.")

# ── RAM allocation for worker count ────────────────────────────────
# RAM_ALLOCATION_MB: explicit memory budget for the transcriber (MB).
#   Empty → auto-detected = total system RAM minus 2 GB reserved for OS.
_RAW_ALLOC_STR = os.environ.get("RAM_ALLOCATION_MB", "").strip()
if _RAW_ALLOC_STR:
    _APP_RAM_MB = int(_RAW_ALLOC_STR)
else:
    _APP_RAM_MB = max(_SYSTEM_RAM_MB - 2048, 2048)  # reserve 2 GB for OS

# WORKER_MEMORY_MB: memory each parallel worker needs (model + audio tensor overhead).
_APP_WORKER_MEM_MB = int(os.environ.get("WORKER_MEMORY_MB", "6144"))
if _APP_WORKER_MEM_MB < 2048:
    _APP_WORKER_MEM_MB = 2048

# Startup info
_alloc_label = f"{_APP_RAM_MB // 1024} GB" if _RAW_ALLOC_STR else f"auto ({_SYSTEM_RAM_MB // 1024} GB system)"
print(f"[INFO] Memory: {_alloc_label} budget, {_APP_WORKER_MEM_MB // 1024} GB/worker")

# ── Chunking for long-form audio (YouTube videos) ────────────────────
CHUNK_DURATION_S = _env_float("CHUNK_DURATION_S", 55 if not _LOW_RAM_MODE else 30)
CHUNK_OVERLAP_S = _env_float("CHUNK_OVERLAP_S", 3)
MAX_AUDIO_BEFORE_SPLIT_S = _env_int("MAX_AUDIO_BEFORE_SPLIT_S", 60)

# Low-RAM overrides for chunking
if _LOW_RAM_MODE:
    CHUNK_DURATION_S = min(CHUNK_DURATION_S, float(_env_float("LOW_RAM_CHUNK_DURATION_S", 25)))
    # On low-RAM systems, smaller max audio before splitting means more chunks
    MAX_AUDIO_BEFORE_SPLIT_S = _env_int("LOW_RAM_MAX_AUDIO_BEFORE_SPLIT_S", 40)

APP_VERSION = _env_str("APP_VERSION", "2.0.0")

# Quality presets → (model_id, chunk_length_s, stride_length_s)
QUALITY_PRESETS = {
    "fast":   (_env_str("FAST_MODEL", "distil-whisper/distil-medium.en"), 20, 3),
    "turbo":  (_env_str("TURBO_MODEL", "kingabzpro/whisper-large-v3-turbo-urdu"), 30, 5),
    "accurate": (_env_str("ACCURATE_MODEL", "openai/whisper-large-v3-turbo"), 30, 5),
}

# ── Parallel / resource control ────────────────────────────────────
_PARALLEL_ENABLED = _env_bool("PARALLEL_ENABLED", True)
_DEFAULT_MAX_WORKERS = _env_int("MAX_WORKERS", 0)   # 0 = auto (based on CPU count)
_WORKERS_PER_CPU = max(_env_int("WORKERS_PER_CPU", 1), 1)

def _resolve_workers(cli_workers: int, device: str | None, total_items: int) -> int:
    """Resolve effective worker count.

    Args:
        cli_workers: Value from --workers CLI flag. 0 = auto (use env/config),
                     >0 = explicit worker count,
                     -1 = disabled (from --no-parallel).
      Precedence (highest → lowest):
      1. --no-parallel (-w -1) or PARALLEL_ENABLED=false → 0 (disabled)
      2. CLI --workers N (explicit override, wins over env/config)
      3. MAX_WORKERS env cap
      4. WORKERS_PER_CPU x cpu_count
      5. config DEFAULT_WORKERS
      6. fallback 1
    """
    # Disabled paths (return immediately)
    if cli_workers == -1:
        return 0
    if not _PARALLEL_ENABLED or total_items <= 1:
        return 0

    # Auto calculation (default when --workers not specified)
    cpu_cores = os.cpu_count() or 4
    max_by_ram = _APP_RAM_MB // _APP_WORKER_MEM_MB
    auto = min(max_by_ram, _WORKERS_PER_CPU * cpu_cores, total_items)
    if _DEFAULT_MAX_WORKERS > 0:
        auto = min(auto, _DEFAULT_MAX_WORKERS)
    cfg_workers = max(DEFAULT_CONFIG.get("workers", 1), 1)

    # CLI explicit override (only when > 0)
    if cli_workers > 0:
        return max(1, min(cli_workers, total_items))

    # Auto path: use auto calc, but ensure at least config default floor
    effective = max(auto, cfg_workers)
    return max(1, effective)

# Default config (user settings override these)
DEFAULT_CONFIG = {
    "model": DEFAULT_MODEL,
    "language": _env_str("DEFAULT_LANGUAGE", "ur"),
    "device": None,           # None = auto-detect
    "format": _env_str("DEFAULT_FORMAT", "txt"),
    "quality": _env_str("QUALITY_PRESET", "turbo"),
    "workers": _env_int("DEFAULT_WORKERS", 1),
    "auto_punctuate": _env_bool("AUTO_PUNCTUATE", True),
    "auto_spell": _env_bool("AUTO_SPELL", True),
    "max_segment_minutes": _env_int("MAX_SEGMENT_MINUTES", 30),
}


# ══════════════════════════════════════════════════════════════════════════════
# CONFIGURATION SYSTEM
# ══════════════════════════════════════════════════════════════════════════════

def _config_path() -> Path:
    # Write config to project root, not src/
    return Path(__file__).parents[1] / CONFIG_FILE


def load_config() -> dict:
    cfg = dict(DEFAULT_CONFIG)
    p = _config_path()
    if p.exists():
        try:
            with open(p, "r", encoding="utf-8") as f:
                cfg.update(json.load(f))
        except (json.JSONDecodeError, IOError):
            pass
    return cfg


def save_config(cfg: dict) -> None:
    with open(_config_path(), "w", encoding="utf-8") as f:
        json.dump(cfg, f, ensure_ascii=False, indent=2)


# ══════════════════════════════════════════════════════════════════════════════
# MODEL SINGLETON — loads once, reused across all files in a run
# ══════════════════════════════════════════════════════════════════════════════

class _ModelCache:
    """Singleton for the Whisper pipeline — loaded once per process."""
    _pipeline = None
    _config_key = None  # (model_id, device_str) to detect mismatches

    @classmethod
    def get(cls, model_id: str, device_config: str | None):
        key = (model_id, str(device_config))
        if cls._pipeline is not None and cls._config_key == key:
            return cls._pipeline
        # Reset cache on mismatch
        cls._pipeline = None
        cls._config_key = None

        import torch
        from transformers import pipeline

        kwargs = {}  # chunk_length_s / stride_length_s are deprecated in transformers >= 5.x

        if device_config:
            if device_config.lower() == "cpu":
                # CPU: use float32 to avoid type mismatch issues
                kwargs["dtype"] = torch.float32
            else:
                # Auto-detect: force CPU (GPU disabled by config)
                kwargs["dtype"] = torch.float32

        # Try loading from cache first (no network needed if cached)
        try:
            cls._pipeline = pipeline(
                "automatic-speech-recognition",
                model=model_id,
                **kwargs,
                local_files_only=True,  # use offline cache when available
            )
        except OSError:
            # Model not cached yet — download it (first run only)
            cls._pipeline = pipeline(
                "automatic-speech-recognition",
                model=model_id,
                **kwargs,
                local_files_only=False,
            )

        cls._config_key = key
        # Clean leftover init params that leak into model.generate()
        if hasattr(cls._pipeline, '_forward_params'):
            cls._pipeline._forward_params.pop('local_files_only', None)
            # Remove duplicate logits processor keys to silence transformers warnings
            for key in ('suppress_tokens', 'bos_token_id', 'decoder_start_token_id'):
                cls._pipeline._forward_params.pop(key, None)
        return cls._pipeline


# ══════════════════════════════════════════════════════════════════════════════
# AUDIO INFO — zero-copy (reads file header only)
# ══════════════════════════════════════════════════════════════════════════════

def get_audio_info(audio_path: Path) -> dict | None:
    """Extract metadata WITHOUT loading audio into memory."""
    try:
        import soundfile as sf
        info = sf.info(str(audio_path))
        return {
            "file": audio_path.name,
            "path": str(audio_path.resolve()),
            "size_mb": round(audio_path.stat().st_size / (1024 * 1024), 2),
            "duration_s": round(info.duration, 2),
            "sample_rate": info.samplerate,
            "channels": info.channels,
            "format": info.format.lower(),
        }
    except Exception:
        # Fallback to librosa for small / unsupported files (loads full file)
        try:
            import librosa
            duration = librosa.get_duration(path=str(audio_path))
            return {
                "file": audio_path.name,
                "path": str(audio_path.resolve()),
                "size_mb": round(audio_path.stat().st_size / (1024 * 1024), 2),
                "duration_s": round(duration, 2),
                "sample_rate": None,
                "channels": None,
                "format": audio_path.suffix.lower(),
            }
        except Exception:
            return None


def format_audio_info(info: dict) -> str:
    lines = [
        f"  File       : {info['file']}",
        f"  Size       : {info['size_mb']} MB",
        f"  Duration   : {_format_duration(info.get('duration_s', 0))}",
    ]
    if info.get("sample_rate"):
        lines.append(f"  Sample Rate: {info['sample_rate']} Hz")
    if info.get("channels"):
        lines.append(f"  Channels   : {info['channels']}")
    if info.get("format"):
        lines.append(f"  Format     : {info['format']}")
    return "\n".join(lines)


def _format_duration(seconds: float) -> str:
    hours, rem = divmod(int(seconds), 3600)
    mins = rem // 60
    secs = rem % 60
    frac = seconds - int(seconds)
    return f"{hours:02d}:{mins:02d}:{secs:02d}.{int(frac * 1000):03d}"


# ══════════════════════════════════════════════════════════════════════════════
# TEXT POST-PROCESSING
# ══════════════════════════════════════════════════════════════════════════════

def normalize_whitespace(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n+", " ", text)
    return text.strip()


# ══════════════════════════════════════════════════════════════════════════════
# CACHE & STATUS HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def compute_file_hash(filepath: Path, chunk_size: int = 8192) -> str:
    h = hashlib.md5()
    with open(filepath, "rb") as f:
        while chunk := f.read(chunk_size):
            h.update(chunk)
    return h.hexdigest()


def load_cache() -> dict:
    p = Path(__file__).parents[1] / CACHE_FILE
    if p.exists():
        try:
            with open(p, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            return {}
    return {}


def save_cache(cache_data: dict) -> None:
    with open(Path(__file__).parents[1] / CACHE_FILE, "w", encoding="utf-8") as f:
        json.dump(cache_data, f, ensure_ascii=False, indent=2)


def load_batch_status() -> dict:
    p = Path(__file__).parents[1] / BATCH_STATUS_FILE
    if p.exists():
        try:
            with open(p, "r", encoding="utf-8") as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError):
            return {}
    return {}


def save_batch_status(status: dict) -> None:
    with open(Path(__file__).parents[1] / BATCH_STATUS_FILE, "w", encoding="utf-8") as f:
        json.dump(status, f, ensure_ascii=False, indent=2)


# ══════════════════════════════════════════════════════════════════════════════
# OUTPUT FORMATTERS
# ══════════════════════════════════════════════════════════════════════════════

def _format_output(text: str, fmt: str, elapsed_s: float, language: str = "ur") -> str | bytes:
    if fmt == "txt":
        return text + "\n"

    if fmt == "json":
        return json.dumps({
            "transcription": text,
            "metadata": {
                "model": DEFAULT_MODEL,
                "language": language,
                "transcription_time_s": round(elapsed_s, 2),
                "generated_at": datetime.now(timezone.utc).isoformat(),
                "version": APP_VERSION,
            },
        }, ensure_ascii=False, indent=2)

    if fmt == "jsonl":
        return json.dumps({
            "transcription": text,
            "metadata": {
                "model": DEFAULT_MODEL,
                "language": language,
                "duration_s": round(elapsed_s, 2),
                "generated_at": datetime.now(timezone.utc).isoformat(),
            },
        }, ensure_ascii=False) + "\n"

    if fmt == "docx":
        from docx import Document
        doc = Document()
        doc.add_paragraph(text)
        doc.add_paragraph(f"Model: {DEFAULT_MODEL} | Language: {language} | Duration: {elapsed_s:.2f}s")
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()





# ══════════════════════════════════════════════════════════════════════════════
# CORE: Single-file transcribe (returns result dict or None)
# ══════════════════════════════════════════════════════════════════════════════

def _resolve_device(device_flag: str | None) -> str | None:
    if device_flag:
        return device_flag.lower()
    # Force CPU only — GPU mode is disabled
    return "cpu"


def _filter_garbage_repetition(text: str, max_repeat: int | None = None) -> str:
    """Remove excessive repetition artifacts from Whisper output.
    
    Whisper can produce garbage like 'دور دور دور دور...' or '\ufeff\ufeff\ufeff'
    when encountering silence, noise, or uncertain audio regions. This function
    reduces runs of the same character/word to at most max_repeat consecutive copies.
    """
    if not text:
        return text

    max_repeat = max_repeat if max_repeat is not None else _env_int("REPEAT_THRESHOLD", 5)

    # Fix 1: Collapse repeated Unicode characters (e.g., '\ufeff' padding, 'ﷺﷺﷺ...')
    def collapse_chars(s: str) -> str:
        result = []
        run_char = ""
        run_count = 0
        for ch in s:
            if ch == run_char and ch != " ":  # skip spaces
                run_count += 1
                if run_count > max_repeat:
                    continue  # skip this char
            else:
                run_char = ch
                run_count = 1
            result.append(ch)
        return "".join(result)

    text = collapse_chars(text)

    # Fix 2: Collapse repeated words (e.g., 'دعا دعا دعا دعا...') 
    # but keep legitimate repetition (sentence structure may intentionally repeat words)
    words = text.split()
    if len(words) <= max_repeat:
        return text

    result_words = []
    run_word = ""
    run_count = 0
    for w in words:
        if w == run_word and w != "":
            run_count += 1
            if run_count > max_repeat:
                continue
        else:
            run_word = w
            run_count = 1
        result_words.append(w)
    return " ".join(result_words)


def transcribe_single(
    audio_path: str,
    *,
    output_path: str | None = None,
    model_id: str = DEFAULT_MODEL,
    language: str = "ur",
    device_flag: str | None = None,
    fmt: str = "txt",
    auto_punctuate: bool = False,
    auto_spell: bool = False,
    quiet: bool = False,
    max_segment_seconds: int | None = None,
) -> dict | None:
    """Transcribe one file. Returns result dict or None on failure."""

    audio_file = Path(audio_path)
    if not audio_file.exists():
        print(f"[ERROR] File not found: {audio_path}")
        return None

    # Resolve output extension
    ext_map = {"txt": ".txt", "json": ".json", "jsonl": ".jsonl", "docx": ".docx"}
    if output_path:
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
    else:
        ext = ext_map.get(fmt, ".txt")
        output_file = audio_file.with_suffix(ext)

    # Header
    if not quiet:
        print(f"\n{'='*60}")
        print(f"  Urdu Speech Recognition  v2")
        print(f"  Model   : {model_id}")
        print(f"  Input   : {audio_file.resolve()}")
        print(f"  Output  : {output_file.resolve()}")
        print(f"{'='*60}\n")

    # Audio info (zero-copy)
    info = get_audio_info(audio_file)
    audio_duration_s = 0.0
    if info:
        audio_duration_s = info.get("duration_s", 0)
        if not quiet:
            print("--- Audio Info ---")
            print(format_audio_info(info))
            print("------------------\n")

    # Resolve device & show message
    resolved_device = _resolve_device(device_flag)
    # GPU mode is disabled; always CPU

    # Load model (singleton — reused across calls in same process)
    try:
        _ModelCache.get(model_id, resolved_device)
    except Exception as e:
        print(f"[ERROR] Failed to load model: {e}")
        return None
    transcriber = _ModelCache._pipeline

    # Force language on the pipeline
    transcriber.model.generation_config.language = language

    print(f"[2/4] Transcribing audio …")
    t_start = time.perf_counter()
    all_segments: list = []
    try:
        from librosa import load as librosa_load

        # Load audio, resample to 16kHz, mono
        audio, sr = librosa_load(str(audio_file), sr=16000, mono=True)
        duration_s = len(audio) / sr

        all_text_parts = []

        if not quiet:
            dur = len(audio) / sr
            print(f"[INFO] {dur:.0f}s audio — transcibed via pipeline")

        # Call the cached pipeline directly for aligned timestamps
        full_result = transcriber(
            audio,
            return_timestamps=True,
            language=language,
        )

        if full_result.get("chunks"):
            for c in full_result["chunks"]:
                ts = c.get("timestamp", (0, None))
                start_s = float(ts[0]) if ts[0] else 0.0
                end_s = float(ts[1]) if ts[1] else start_s + 1.0
                all_segments.append({
                    "start": round(start_s, 3),
                    "end": round(end_s, 3),
                    "text": c.get("text", "").strip(),
                })
            all_text_parts = [full_result.get("text", "")]
        else:
            # Pipeline returned text but no chunks — create estimated segments
            txt = full_result.get("text", "") or ""
            if txt:
                all_text_parts = [txt]
            if _LOW_RAM_MODE or resolved_device == "cpu":
                import gc as _gc
                _gc.collect()

    except Exception as e:
        print(f"[ERROR] Transcription failed: {e}")
        import traceback
        traceback.print_exc()
        return None

    elapsed_s = time.perf_counter() - t_start
    text = " ".join(all_text_parts).strip()

    # Filter out excessive repetition artifacts (silence/garbage produced by Whisper)
    text = _filter_garbage_repetition(text)

    if not text:
        print("[WARNING] Transcription returned empty text.")
        return {"text": "", "output_path": str(output_file), "skipped": False, "duration_s": round(elapsed_s, 2)}

    # Post-process
    if auto_punctuate:
        text = normalize_whitespace(text)

    # Apply Urdu spell/word correction if requested (cached corrector to avoid re-loading)
    _corrector = getattr(transcribe_single, "_cached_corrector", None)
    if _corrector is None:
        from urdu_correction import UrduSpellCorrector  # type: ignore
        _corrector = UrduSpellCorrector()
        transcribe_single._cached_corrector = _corrector
    if auto_spell:
        text, correction_stats = _corrector.correct(text, level="full")

    # Speed ratio
    speed_ratio = audio_duration_s / elapsed_s if elapsed_s > 0 else None
    if not quiet and HAS_TQDM and audio_duration_s > 0 and elapsed_s > 0.5:
        arrow = ">=" if speed_ratio >= 1 else "<"
        print(f"      Speed: {speed_ratio:.2f}x real-time  ({arrow})")

    # Write output
    content = _format_output(text, fmt, elapsed_s, language=language)
    print(f"[3/4] Saving output …")
    if isinstance(content, bytes):
        with open(output_file, "wb") as f:
            f.write(content)
    else:
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(content)

    if not quiet:
        print(f"[OK]  Transcription saved to: {output_file.resolve()}")
        print(f"      Done in {elapsed_s:.1f}s ({audio_duration_s:.0f}s audio)")

    return {"text": text, "segments": all_segments if all_segments else None, "output_path": str(output_file), "skipped": False, "duration_s": round(elapsed_s, 2)}


# ══════════════════════════════════════════════════════════════════════════════
# CHUNKED LONG-FORM TRANSCRIPTION
# ══════════════════════════════════════════════════════════════════════════════

def _worker_transcribe_chunk(args: tuple) -> dict:
    """Worker function — runs in subprocess. Transcribes ONE chunk file."""
    audio_path, model_id, language, device_flag, quiet = args
    # Import inside worker to avoid pickling issues with shared pipeline
    sys.path.insert(0, str(Path(__file__).parent))
    from transcribe import transcribe_single  # type: ignore
    result = transcribe_single(
        audio_path=audio_path,
        output_path=None,
        model_id=model_id,
        language=language,
        device_flag=device_flag,
        fmt="txt",
        auto_punctuate=False,  # will be applied after merge
        auto_spell=False,
        quiet=quiet,
    )
    text = result["text"] if result else ""
    segments = result.get("segments") if result else None
    return {"path": audio_path, "text": text, "segments": segments, "ok": result is not None}


def transcribe_in_chunks(
    audio_path: str,
    *,
    output_path: str | None = None,
    model_id: str = DEFAULT_MODEL,
    language: str = "ur",
    device_flag: str | None = None,
    fmt: str = "txt",
    auto_spell: bool = False,
    max_audio_before_split_s: int = MAX_AUDIO_BEFORE_SPLIT_S,
) -> dict | None:
    """Split long audio via ffmpeg, transcribe each chunk on CPU in parallel, merge results.

    This is the preferred path for YouTube videos > 60s. Each chunk gets a fresh
    Whisper context window (no cross-chunk context degradation).

    Args:
        audio_path: Path to input audio file.
        output_path: Where to write output.
        model_id, language, device_flag, fmt, auto_spell: forwarded to transcribe_single.
        max_audio_before_split_s: Only split if audio exceeds this duration.

    Returns:
        Result dict (same shape as transcribe_single) or None on failure.
    """
    audio_file = Path(audio_path)
    if not audio_file.exists():
        print(f"[ERROR] File not found: {audio_path}")
        return None

    # ── Step 0: resolve output ─────────────────────────────────────────────
    ext_map = {"txt": ".txt", "json": ".json", "jsonl": ".jsonl", "docx": ".docx"}
    if output_path:
        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)
    else:
        ext = ext_map.get(fmt, ".txt")
        output_file = audio_file.with_suffix(ext)

    # ── Step 1: probe duration ─────────────────────────────────────────────
    info = get_audio_info(audio_file)
    duration_s = info["duration_s"] if info else 0.0

    # If short enough, skip splitting — use normal transcribe_single
    if duration_s <= max_audio_before_split_s:
        return transcribe_single(
            audio_path=audio_path, output_path=output_path,
            model_id=model_id, language=language,
            device_flag=device_flag, fmt=fmt,
            auto_punctuate=True, auto_spell=auto_spell, quiet=False,
            max_segment_seconds=_env_int("SHORT_AUDIO_MAX_SEGMENT_TXT", 120),  # can safely use larger chunks for short audio
        )

    print(f"\n{'='*60}")
    print(f"  Long-Form Transcription  v3")
    print(f"  Model   : {model_id}")
    print(f"  Input   : {audio_file.resolve()}")
    print(f"  Duration: {duration_s:.0f}s (auto-split enabled)")
    print(f"  Output  : {output_file.resolve()}")
    print(f"{'='*60}\n")

    # ── Step 2: ffmpeg split ───────────────────────────────────────────────
    sys.path.insert(0, str(Path(__file__).parent))
    from utils.audio_splitter import split_audio, calculate_chunk_start_times  # type: ignore

    chunks_dir = str(audio_file.parent / f"{audio_file.stem}_chunks")
    chunk_paths = split_audio(
        audio_path,
        chunk_duration_s=CHUNK_DURATION_S,
        overlap_s=CHUNK_OVERLAP_S,
        output_dir=chunks_dir,
        prefix="c",
    )
    if not chunk_paths:
        print("[ERROR] No chunks created — aborting.")
        return None

    total_chunks = len(chunk_paths)

    # ── Step 3: CPU transcribe (sequential on low-RAM, parallel otherwise) ───
    # Each worker loads the model once. Workers = RAM_ALLOCATION_MB / WORKER_MEMORY_MB,
    # capped by MAX_WORKERS, WORKERS_PER_CPU, and total_chunks.
    if _LOW_RAM_MODE or not _PARALLEL_ENABLED:
        print(f"[INFO] Using sequential processing (memory-safe mode) — {total_chunks} chunk(s)")
        max_workers = 1
    else:
        effective_workers = _resolve_workers(cli_workers=0, device=device_flag, total_items=total_chunks)
        max_workers = effective_workers if effective_workers > 1 else 1

    print(f"[2/4] Transcribing {total_chunks} chunk(s) with {max_workers} worker(s) …")

    t_start = time.perf_counter()
    chunk_texts: list[str | None] = [None] * total_chunks
    chunk_segments: list[list | None] = [None] * total_chunks
    chunk_ok = [False] * total_chunks

    worker_args_list = [(p, model_id, language, device_flag, False) for p in chunk_paths]

    # Create executor and store reference on the module so Ctrl+C handler can shut it down
    _executor_ref = ProcessPoolExecutor(max_workers=max(1, max_workers))
    completed = 0
    failed = 0
    try:
        futures = {_executor_ref.submit(_worker_transcribe_chunk, a): i for i, a in enumerate(worker_args_list)}
        pbar = tqdm(total=total_chunks, desc="Chunks", disable=not HAS_TQDM)
        for future in as_completed(futures):
            i = futures[future]
            completed += 1
            try:
                result = future.result()
                if result.get("ok"):
                    chunk_texts[i] = result["text"]
                    chunk_segments[i] = result.get("segments")
                    chunk_ok[i] = True
                    pbar.update(1)
                else:
                    failed += 1
                    print(f"\n[WARNING] Chunk {i} ({Path(worker_args_list[i][0]).name}) transcription failed — text will be empty.")
            except Exception as e:
                failed += 1
                err_msg = str(e).lower()
                if any(kw in err_msg for kw in ("memory", "oom", "paging file", "out of memory", "allocation failed")):
                    print(f"\n[ERROR] Worker for chunk {i} crashed — likely OOM/memory error.")
                    print(f"        Reduce workers by setting WORKER_MEMORY_MB higher, or use --no-parallel / -w 1.")
                else:
                    print(f"\n[ERROR] Worker for chunk {i} crashed: {e}")
            pbar.refresh()
            # Print explicit progress every N chunks (or on last) so user always sees advancement
            if completed % max(1, total_chunks // 5) == 0 or completed == total_chunks:
                remaining = total_chunks - completed - failed
                print(f"\r[Progress] {completed}/{total_chunks} done ({failed} failed, {remaining} remaining…)", end="", flush=True)
    except KeyboardInterrupt:
        pbar.close()
        raise
    finally:
        _executor_ref.shutdown(wait=False, cancel_futures=True)
        _executor_ref = None  # clear reference after shutdown
    pbar.close()

    # ── Step 4: merge ──────────────────────────────────────────────────────
    elapsed_s = time.perf_counter() - t_start

    if fmt == "json":
        # JSON uses plain text with metadata — use merge_text
        cleaned = [t for t in chunk_texts if t]
        merged_text = " ".join(cleaned) if cleaned else ""
    else:
        from utils.audio_splitter import merge_text  # type: ignore
        merged_text = merge_text([t if t else "" for t in chunk_texts])

    # ── Step 5: post-process & save ────────────────────────────────────────
    text = re.sub(r"[ \t]+", " ", merged_text).strip()

    # Free chunk data (no longer needed)
    del chunk_texts
    del chunk_ok
    if _LOW_RAM_MODE:
        import gc as _gc2
        _gc2.collect()

    if auto_spell:
        sys.path.insert(0, str(Path(__file__).parent))
        _corr = getattr(transcribe_in_chunks, "_cached_corrector", None)
        if _corr is None:
            from urdu_correction import UrduSpellCorrector  # type: ignore
            _corr = UrduSpellCorrector()
            transcribe_in_chunks._cached_corrector = _corr
        text, _ = _corr.correct(text, level="full")

    content = _format_output(text, fmt, elapsed_s, language=language)
    print(f"[3/4] Saving output …")
    if isinstance(content, bytes):
        with open(output_file, "wb") as f:
            f.write(content)
    else:
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(content)

    # Clean up chunk directory
    try:
        import shutil
        shutil.rmtree(chunks_dir, ignore_errors=True)
    except Exception:
        pass

    speed_ratio = duration_s / elapsed_s if elapsed_s > 0 else None
    print(f"[OK]  Transcription saved to: {output_file.resolve()}")
    print(f"      Done in {elapsed_s:.1f}s ({duration_s:.0f}s audio, {total_chunks} chunks)")
    if speed_ratio:
        arrow = ">=" if speed_ratio >= 1 else "<"
        print(f"      Speed: {speed_ratio:.2f}x real-time  ({arrow})")

    return {
        "text": text,
        "output_path": str(output_file),
        "skipped": False,
        "duration_s": round(elapsed_s, 2),
    }


# ══════════════════════════════════════════════════════════════════════════════
# CORE: Batch processing (with model reuse, resume, parallel)
# ══════════════════════════════════════════════════════════════════════════════

def _worker(args: tuple) -> dict:
    """Worker for parallel batch — runs in subprocess."""
    audio_path, output_path, model_id, language, device_flag, fmt, auto_punctuate = args
    result = transcribe_single(
        audio_path=audio_path, output_path=output_path,
        model_id=model_id, language=language,
        device_flag=device_flag, fmt=fmt,
        auto_punctuate=auto_punctuate, quiet=True,
    )
    return {"path": audio_path, "result": result}


def run_batch(
    directory: str,
    *,
    model_id: str = DEFAULT_MODEL,
    language: str = "ur",
    device_flag: str | None = None,
    fmt: str = "txt",
    use_cache: bool = True,
    workers: int = 1,
    auto_punctuate: bool = False,
    auto_spell: bool = False,
) -> dict:
    """Process all audio files in a directory. Returns stats."""

    dir_path = Path(directory)
    if not dir_path.is_dir():
        print(f"[ERROR] Not a directory: {directory}")
        return {}

    # Recursively find supported files
    audio_files = sorted([
        f for f in dir_path.rglob("*")
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
    ])
    if not audio_files:
        print(f"[INFO] No supported audio files found in: {directory}")
        return {}

    total = len(audio_files)

    # Resolve effective worker count
    effective_workers = _resolve_workers(workers, device_flag, total)
    parallel = effective_workers > 1

    print(f"\n{'='*60}")
    print(f"  Batch Transcription  v2")
    print(f"  Directory : {dir_path.resolve()}")
    print(f"  Files     : {total} audio file(s)")
    print(f"  Workers   : {'parallel (auto)' if parallel else 'sequential'}")
    if effective_workers > 1:
        print(f"  Memory  : {_APP_RAM_MB // 1024} GB budget, {_APP_WORKER_MEM_MB // 1024} GB/worker → {effective_workers} workers")
        cpu_cores = os.cpu_count() or 4
        print(f"  Config    : max_workers={_DEFAULT_MAX_WORKERS or 'auto'}, workers_per_cpu={_WORKERS_PER_CPU}")
    if not _PARALLEL_ENABLED:
        print(f"  Parallel  : disabled (PARALLEL_ENABLED=false)")
    print(f"  Cache     : {'enabled' if use_cache else 'disabled'}")
    print(f"{'='*60}\n")

    # Resume: check status file
    batch_status = load_batch_status()
    dir_key = str(dir_path.resolve())
    if batch_status.get("directory") != dir_key:
        batch_status = {}  # new directory — clear old

    succeeded = skipped = failed = 0
    overall_start = time.perf_counter()

    work_items = []
    for af in audio_files:
        key = str(af.resolve())
        if batch_status.get(key, {}).get("status") == "done":
            skipped += 1
            continue
        work_items.append((str(af), None, model_id, language, device_flag, fmt, auto_punctuate))

    if not work_items:
        print("[INFO] All files already completed (resume from status file).")
        return {"total": total, "succeeded": succeeded, "skipped": skipped, "failed": failed}

    # Execute
    if parallel:
        print(f"[INFO] Using {effective_workers} parallel workers …")
        bar = tqdm(total=len(work_items), desc="Processing", unit="file", disable=not HAS_TQDM)
        _executor_ref = ProcessPoolExecutor(max_workers=min(effective_workers, len(work_items)))
        try:
            futures = {_executor_ref.submit(_worker, item): item for item in work_items}
            for future in as_completed(futures):
                try:
                    file_result = future.result()
                    key = file_result["path"]
                    r = file_result["result"]
                    if r is None:
                        failed += 1
                    elif r.get("skipped"):
                        skipped += 1
                    else:
                        succeeded += 1
                        batch_status[key] = {"status": "done", "completed_at": datetime.now(timezone.utc).isoformat()}
                        save_batch_status(batch_status)
                    bar.update(1)
                except Exception as e:
                    failed += 1
                    print(f"\n[ERROR] Worker crash: {e}")
        except KeyboardInterrupt:
            bar.close()
            raise
        finally:
            _executor_ref.shutdown(wait=False, cancel_futures=True)
            _executor_ref = None  # clear reference after shutdown
        bar.close()
    else:
        print(f"[INFO] Processing {len(work_items)} files sequentially …\n")
        for i, args in enumerate(work_items, 1):
            af_path = Path(args[0])
            t_file = time.perf_counter()

            if HAS_TQDM and not any("--quiet" in sys.argv or "-q" in sys.argv for _ in (1,)):
                print(f"\n[{i}/{len(work_items)}] {af_path.name}")

            result = transcribe_single(
                audio_path=args[0], output_path=None,
                model_id=model_id, language=language,
                device_flag=device_flag, fmt=fmt,
                auto_spell=False,
                quiet=False,
            )

            key = str(af_path.resolve())
            if result is None:
                failed += 1
            elif result.get("skipped"):
                skipped += 1
            else:
                succeeded += 1
                batch_status[key] = {"status": "done", "completed_at": datetime.now(timezone.utc).isoformat()}

            save_batch_status(batch_status)

    total_elapsed = time.perf_counter() - overall_start
    batch_status["directory"] = dir_key
    batch_status["last_run"] = datetime.now(timezone.utc).isoformat()
    save_batch_status(batch_status)

    print(f"\n{'='*60}")
    print(f"  Batch Complete!")
    print(f"  Succeeded : {succeeded}/{total}")
    if skipped:
        print(f"  Skipped   : {skipped} (cached or completed)")
    if failed:
        print(f"  Failed    : {failed}")
    print(f"  Wall time : {total_elapsed:.1f}s")
    if succeeded > 0:
        print(f"  Avg per file: {(total_elapsed / succeeded):.1f}s")
    print(f"{'='*60}\n")

    return {"total": total, "succeeded": succeeded, "skipped": skipped, "failed": failed, "total_time_s": round(total_elapsed, 2)}


# ══════════════════════════════════════════════════════════════════════════════
# DRY-RUN HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _dry_run_single(audio_path: str, config: dict) -> None:
    f = Path(audio_path)
    info = get_audio_info(f)
    fmt = config.get("format", "txt")
    ext_map = {"txt": ".txt", "json": ".json", "jsonl": ".jsonl"}

    print(f"\n{'='*60}")
    print(f"  Dry Run Preview")
    print(f"{'='*60}\n")

    if info:
        print("--- Audio Info ---")
        print(format_audio_info(info))
        print("------------------\n")

    print(f"  Model         : {config.get('model', DEFAULT_MODEL)}")
    print(f"  Quality       : {config.get('quality', 'turbo')}")
    print(f"  Language      : {config['language']}")
    print(f"  Device        : {config['device'] or '(auto-detect)'}")
    print(f"  Format        : {fmt}")
    print(f"  Output file   : {f.with_suffix(ext_map.get(fmt, '.txt')).resolve()}")
    print(f"  Auto-punctuate: {'yes' if config.get('auto_punctuate') else 'no'}")
    cache = load_cache()
    hit = str(f.resolve()) in cache and cache[str(f.resolve())].get("model") == config.get("model")
    print(f"  Cache status  : {'found (would skip)'} if {hit} else 'not found (would process)')")
    print(f"\n{'='*60}")


def _dry_run_batch(directory: str, config: dict) -> None:
    dir_path = Path(directory)
    audio_files = sorted([f for f in dir_path.rglob("*") if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS])
    print(f"\n{'='*60}")
    print(f"  Dry Run — Batch Preview")
    print(f"{'='*60}\n")
    print(f"  Directory  : {dir_path.resolve()}")
    print(f"  Files found: {len(audio_files)}")
    print(f"  Model      : {config.get('model', DEFAULT_MODEL)}")
    print(f"  Format     : {config.get('format', 'txt')}\n")

    total_dur = 0.0
    cache = load_cache()
    for af in audio_files[:20]:
        info = get_audio_info(af)
        dur = info["duration_s"] if info and isinstance(info.get("duration_s"), (int, float)) else "?"
        if isinstance(dur, (int, float)):
            total_dur += dur
        hit = str(af.resolve()) in cache and cache[str(af.resolve())].get("model") == config.get("model")
        status = "cached" if hit else "would process"
        print(f"  {status:>12} : {af.name:<40s} {dur}s")

    if len(audio_files) > 20:
        print(f"  ... and {len(audio_files) - 20} more files")
    print(f"\n{'='*60}")
    print(f"  Total audio duration : {_format_duration(total_dur)}")
    print(f"  Estimated transcribe : ~{_format_duration(total_dur / 5)} (at ~5x speed)")
    print(f"{'='*60}\n")


# ══════════════════════════════════════════════════════════════════════════════
# CLI — subcommands: trans, batch, config
# ══════════════════════════════════════════════════════════════════════════════

def _build_trans_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(add_help=True, prog="transcribe.py trans")
    p.add_argument("audio", help="Audio file path (mp3, wav, flac, ogg, m4a, mp4, webm)")
    p.add_argument("--output", "-o", default=None, help="Output file path")
    p.add_argument("--format", "-f", choices=["txt", "json", "jsonl", "docx"], default="txt", help="Output format (default: txt)")
    p.add_argument("--language", "-l", default=None, help="Language code (e.g., ur, en, hi, ar)")
    p.add_argument("--model", "-m", default=None, help="HuggingFace model ID")
    p.add_argument("--device", choices=["cpu", "gpu"], default=None, help="Force device (CPU only; GPU mode is disabled)")
    p.add_argument("--quality", choices=["fast", "turbo", "accurate"], default=None, help="Speed/quality preset")
    p.add_argument("--no-cache", action="store_true", help="Disable cache")
    p.add_argument("--quiet", "-q", action="store_true", help="Minimal output (JSON to stdout)")
    p.add_argument("--dry-run", action="store_true", help="Preview without processing")
    p.add_argument("--info", action="store_true", help="Show audio info only")
    p.add_argument("--spell", "-s", action="store_true", help="Apply Urdu spell/word correction (NEW v2 feature)")
    p.add_argument("--chunk", type=float, default=0, help="Split audio into chunks of N seconds (e.g. 55). Use 0 to disable.")
    return p


def _build_batch_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(add_help=True, prog="transcribe.py batch")
    p.add_argument("directory", nargs="?", default=None, help="Directory to scan for audio files")
    p.add_argument("--format", "-f", choices=["txt", "json", "jsonl", "docx"], default="txt", help="Output format (default: txt)")
    p.add_argument("--language", "-l", default=None, help="Language code")
    p.add_argument("--model", "-m", default=None, help="HuggingFace model ID")
    p.add_argument("--device", choices=["cpu", "gpu"], default=None, help="Force device (CPU only; GPU mode is disabled)")
    p.add_argument("--workers", "-w", type=int, default=0, help="Parallel workers for batch (CPU only; 0=auto, -w 1=disable)")
    p.add_argument("--parallel", dest="parallel", action="store_true", help="Force parallel processing (overrides PARALLEL_ENABLED=false)")
    p.add_argument("--no-parallel", dest="no_parallel", action="store_true", help="Disable parallel processing")
    p.add_argument("--no-cache", action="store_true", help="Disable cache")
    p.add_argument("--dry-run-batch", dest="dry_run_batch", metavar="DIR", help="Preview batch without processing")
    p.add_argument("--spell", "-s", action="store_true", help="Apply Urdu spell/word correction (NEW v2 feature)")
    return p


def _build_config_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(add_help=False)
    sub = p.add_subparsers(dest="action")
    sub.add_parser("show", help="Show current configuration")
    s = sub.add_parser("set", help="Set config values")
    s.add_argument("kv_pairs", nargs="+", help="key=value pairs (e.g., model=my-model language=en)")
    return p


def _cli_main():
    """Entry point — no ambiguity between positional args and subcommands."""

    if len(sys.argv) < 2 or sys.argv[1] in ("--help", "-h"):
        _print_usage()
        sys.exit(0)

    cmd = sys.argv[1].lower()

    # ── config subcommand ──────────────────────────────────────────────────
    if cmd == "config":
        cfg_parser = _build_config_parser()
        args, extra = cfg_parser.parse_known_args(sys.argv[2:])
        cfg = load_config()

        if args.action == "show":
            print("\n=== Current Configuration ===\n")
            for k, v in cfg.items():
                print(f"  {k.replace('_', '-'):<20s}: {v}")
            print("\n")
        elif args.action == "set":
            for kv in args.kv_pairs:
                if "=" not in kv:
                    print(f"[ERROR] Use key=value format. Got: '{kv}'")
                    sys.exit(1)
                key, val = kv.split("=", 1)
                key = key.strip().replace("-", "_")
                val = val.strip()
                if val.lower() == "true":
                    val = True
                elif val.lower() == "false":
                    val = False
                elif val.isdigit():
                    val = int(val)
                cfg[key] = val
            save_config(cfg)
            print(f"[OK] Configuration saved ({len(args.kv_pairs)} setting(s)).")
        else:
            cfg_parser.print_help()
        return

    # ── trans subcommand (single file) ─────────────────────────────────────
    if cmd == "trans":
        if len(sys.argv) > 2 and sys.argv[2] in ("--help", "-h"):
            parser = _build_trans_parser()
            parser.print_help()
            return
        parser = _build_trans_parser()
        args = parser.parse_args(sys.argv[2:])
        cfg = load_config()

        if args.info:
            af = Path(args.audio)
            if not af.exists():
                print(f"[ERROR] File not found: {args.audio}")
                sys.exit(1)
            info = get_audio_info(af)
            if info:
                print("--- Audio Info ---")
                print(format_audio_info(info))
                print("------------------")
            else:
                print("[WARNING] Could not read audio metadata.")
            return

        if args.dry_run:
            _dry_run_single(args.audio, cfg)
            return

        # Resolve config → CLI precedence
        quality = args.quality or cfg.get("quality", "turbo")
        model_id = args.model or cfg.get("model", DEFAULT_MODEL)
        if quality != "turbo" and quality in QUALITY_PRESETS:
            model_id = QUALITY_PRESETS[quality][0]

        # Resolve language: CLI > config > default
        lang = args.language or cfg.get("language", "ur")

        # Auto-detect long audio: if --chunk > 0 (or default chunking enabled), split
        use_chunked = args.chunk > 0
        auto_split = False

        if use_chunked:
            max_audio = int(args.chunk * 1.2)  # split files longer than ~chunk_duration
        else:
            # Auto-detect: probe file duration, split if > 60s for txt/json
            af = Path(args.audio)
            if af.exists():
                info = get_audio_info(af)
                dur = info["duration_s"] if info else 0
                if dur > MAX_AUDIO_BEFORE_SPLIT_S:
                    auto_split = True
                    use_chunked = True

        if use_chunked:
            result = transcribe_in_chunks(
                audio_path=args.audio,
                output_path=args.output,
                model_id=model_id,
                language=lang,
                device_flag=args.device or cfg.get("device"),
                fmt=args.format,
                auto_spell=args.spell,
                max_audio_before_split_s=int(args.chunk * 1.2) if use_chunked and not auto_split else MAX_AUDIO_BEFORE_SPLIT_S,
            )
        else:
            result = transcribe_single(
                audio_path=args.audio,
                output_path=args.output,
                model_id=model_id,
                language=lang,
                device_flag=args.device or cfg.get("device"),
                fmt=args.format,
                auto_punctuate=cfg.get("auto_punctuate", True),
                quiet=args.quiet,
            )
        if result is None:
            sys.exit(1)

        # Quiet mode → JSON to stdout
        if args.quiet:
            print(json.dumps({
                "status": "ok" if not result.get("skipped") else "skipped",
                "text": result.get("text", ""),
                "output_path": result.get("output_path"),
                "duration_s": result.get("duration_s"),
            }))
        return

    # ── batch subcommand ───────────────────────────────────────────────────
    if cmd == "batch":
        if len(sys.argv) > 2 and sys.argv[2] in ("--help", "-h"):
            parser = _build_batch_parser()
            parser.print_help()
            return
        parser = _build_batch_parser()
        args = parser.parse_args(sys.argv[2:])
        cfg = load_config()

        if args.dry_run_batch:
            _dry_run_batch(args.dry_run_batch, cfg)
            return

        if not args.directory:
            print("[ERROR] Provide a directory path.")
            parser.print_help()
            sys.exit(1)

        quality = cfg.get("quality", "turbo")
        model_id = args.model or cfg.get("model", DEFAULT_MODEL)
        if quality != "turbo" and quality in QUALITY_PRESETS:
            model_id = QUALITY_PRESETS[quality][0]

        # Resolve parallel flag: CLI > env > config
        workers_arg = args.workers  # default 0 = auto
        if args.no_parallel:
            workers_arg = -1  # sentinel for disabled
        elif args.parallel:
            workers_arg = max(workers_arg, 2)  # ensure > 1 to force parallel

        result = run_batch(
            directory=args.directory,
            model_id=model_id,
            language=args.language or cfg.get("language", "ur"),
            device_flag=args.device or cfg.get("device"),
            fmt=args.format,
            use_cache=not args.no_cache,
            workers=workers_arg,
            auto_punctuate=cfg.get("auto_punctuate", True),
            auto_spell=args.spell,  # NEW: Pass spell flag through
        )
        sys.exit(1 if result.get("failed", 0) > 0 else 0)

    # ── Unknown command ────────────────────────────────────────────────────
    print(f"[ERROR] Unknown command: '{sys.argv[1]}'.")
    _print_usage()
    sys.exit(1)


def _print_usage():
    print("""
Urdu Speech-to-Text Transcription Tool  v2
============================================

USAGE:
  python transcribe.py trans <audio_file> [options]   Single-file transcription
  python transcribe.py batch <directory> [options]    Batch process all files
  python transcribe.py config show                    Show current configuration
  python transcribe.py config set key=value           Update configuration

SINGLE-FILE OPTIONS (trans):
  --format f        txt | json | jsonl | docx   (default: txt)
  --language l      ur | en | hi | ar             (default: ur from config or 'ur')
  --model m         Custom HuggingFace model ID
  --device          cpu                        (CPU only; GPU mode is disabled)
  --quality         fast | turbo | accurate        (turbo = default; changes model)
  --output o        Save to specific file path
  --quiet -q        JSON output only (for scripts/APIs)
  --dry-run         Preview without processing
  --info            Show audio info, do not transcribe
  --chunk SEC       Split audio into chunks of SEC seconds (default: auto >60s, use 0 to disable)

BATCH OPTIONS (batch):
  --format f        txt | json | jsonl | docx   (default: txt)
  --language l      ur | en | hi | ar                 (default: from config)
  --model m         Custom HuggingFace model ID
  --device          cpu                         (CPU only; GPU mode is disabled)
  --workers -w N    Parallel workers (CPU only; 0=auto, -w 1=disable)
  --parallel        Force parallel (overrides PARALLEL_ENABLED=false)
  --no-parallel     Disable parallel processing
  --output o        Base output path (each file gets its own file)
  --no-cache        Disable cache for fresh processing
  --dry-run-batch DIR  Preview without processing

CONFIG OPTIONS (config):
  show              Display current configuration
  set key=value     Update one or more settings (e.g. model=my-model language=en)

EXAMPLES:
  python transcribe.py trans speech.mp3
  python transcribe.py trans lecture.mp4 --format json
  python transcribe.py trans audio.wav --language en
  python transcribe.py trans audio.mp3 --quality fast
  python transcribe.py batch ./podcasts/
  python transcribe.py batch ./lectures/ -w 4
  python transcribe.py batch ./audio/ --no-parallel
  python transcribe.py config show
  python transcribe.py config set language=en format=json
""")


# ══════════════════════════════════════════════════════════════════════════════
# Graceful shutdown on Ctrl+C
# ══════════════════════════════════════════════════════════════════════════════

_executor_ref = None  # track active ProcessPoolExecutor for cleanup
_interrupt_registered = False

def _shutdown_handler(sig, frame):
    """Handle Ctrl+C by shutting down the active executor and cleaning up."""
    global _executor_ref
    print("\n\n[INTERRUPT] Shutting down…")
    if _executor_ref is not None:
        try:
            _executor_ref.shutdown(wait=False, cancel_futures=True)
        except Exception:
            pass
        _executor_ref = None
    # Clean up leftover chunk directories
    import glob as _glob
    import shutil as _shutil
    for d in _glob.glob(str(Path(__file__).parents[1] / "*" / "*_chunks")):
        try:
            _shutil.rmtree(d, ignore_errors=True)
        except Exception:
            pass
    sys.exit(130)  # standard exit code for SIGINT

def _register_interrupt_handler():
    global _interrupt_registered
    if _interrupt_registered:
        return
    _interrupt_registered = True
    import signal
    try:
        signal.signal(signal.SIGINT, _shutdown_handler)
        signal.signal(signal.SIGTERM, _shutdown_handler)
    except (OSError, ValueError):
        pass  # Windows or signal in main thread issues

# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    _register_interrupt_handler()
    _cli_main()

# Also register on import so youtube_transcribe.py gets Ctrl+C protection
_register_interrupt_handler()
